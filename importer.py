import docx
import sqlite3
import re
import os
from datetime import datetime

class EohIndexImporter:
    def __init__(self, db_path='adidam_recordings.db'):
        self.db_path = db_path
        
    def import_from_docx(self, docx_path):
        """Import data from the EOH Index Word document"""
        if not os.path.exists(docx_path):
            print(f"Error: File not found: {docx_path}")
            return False
            
        try:
            # Open the database connection
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            # Open the Word document
            doc = docx.Document(docx_path)
            
            print(f"Document opened: {docx_path}")
            print(f"Total paragraphs: {len(doc.paragraphs)}")
            
            # Start a transaction
            conn.execute('BEGIN TRANSACTION')
            
            current_book = None
            book_id = None
            
            # Process each paragraph
            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if not text:
                    continue
                
                # Print a progress indicator
                if i % 100 == 0:
                    print(f"Processing paragraph {i}...")
                
                # Parse the paragraph
                self.process_paragraph(text, cursor)
            
            # Commit the transaction
            conn.commit()
            print("Import completed successfully")
            return True
            
        except Exception as e:
            if conn:
                conn.rollback()
            print(f"Error during import: {str(e)}")
            return False
        finally:
            if conn:
                conn.close()
    
    def process_paragraph(self, text, cursor):
        """Process a paragraph from the document"""
        # Check if it's a book title (starts with *[ and ends with ]*)
        if text.startswith("*[") and text.endswith("]*"):
            self.process_book_title(text, cursor)
        
        # Check if it's an essay entry (contains ** followed by numbers)
        elif "**" in text:
            self.process_essay_entry(text, cursor)
    
    def process_book_title(self, text, cursor):
        """Process a book title paragraph"""
        # Extract book title, removing formatting markers
        book_title = text.replace("*[", "").replace("]*", "").replace("{.smallcaps}", "").strip()
        
        # Check if this book already exists
        cursor.execute("SELECT id FROM books WHERE title = ?", (book_title,))
        existing_book = cursor.fetchone()
        
        if existing_book:
            self.current_book_id = existing_book[0]
            self.current_book_title = book_title
            print(f"Found existing book: {book_title} (ID: {self.current_book_id})")
        else:
            # Insert new book
            cursor.execute(
                "INSERT INTO books (title, display_order) VALUES (?, ?)",
                (book_title, len(book_title))  # Use title length as simple display order
            )
            self.current_book_id = cursor.lastrowid
            self.current_book_title = book_title
            print(f"Added new book: {book_title} (ID: {self.current_book_id})")
    
    def process_essay_entry(self, text, cursor):
        """Process an essay entry paragraph"""
        if not hasattr(self, 'current_book_id'):
            print("Warning: Essay entry found before book title. Skipping.")
            return
        
        # Use regex to extract essay number and title
        # Pattern: **number** title
        match = re.match(r'\s*\*\*([^*]+)\*\*\s+(.*)', text)
        if match:
            essay_number = match.group(1).strip()
            essay_title = match.group(2).strip()
            
            # Clean up the title (remove formatting)
            essay_title = re.sub(r'\[([^\]]+)\]\.[\w]+', r'\1', essay_title)
            essay_title = re.sub(r'\{\.[\w]+\}', '', essay_title)
            essay_title = essay_title.replace("[", "").replace("]", "").replace("{.underline}", "")
            
            # Handle multiple recording numbers for the same essay
            essay_numbers = [num.strip() for num in essay_number.split(',')]
            
            for num in essay_numbers:
                # Check if this essay already exists
                cursor.execute(
                    "SELECT id FROM essays WHERE title = ? AND book_id = ? AND essay_number = ?", 
                    (essay_title, self.current_book_id, num)
                )
                existing_essay = cursor.fetchone()
                
                if existing_essay:
                    essay_id = existing_essay[0]
                    print(f"  Found existing essay: {essay_title} (ID: {essay_id}, Number: {num})")
                else:
                    # Insert new essay
                    cursor.execute(
                        """
                        INSERT INTO essays 
                        (title, book_id, essay_number, display_order) 
                        VALUES (?, ?, ?, ?)
                        """,
                        (essay_title, self.current_book_id, num, int(num) if num.isdigit() else 0)
                    )
                    essay_id = cursor.lastrowid
                    print(f"  Added new essay: {essay_title} (ID: {essay_id}, Number: {num})")
                    
                    # For each essay, create a placeholder recording
                    cursor.execute(
                        """
                        INSERT INTO recordings 
                        (essay_id, title, date_added) 
                        VALUES (?, ?, ?)
                        """,
                        (essay_id, f"Recording of {essay_title}", datetime.now().isoformat())
                    )

def main():
    print("Adidam EOH Index Importer")
    print("=========================")
    
    # Ask for document file
    docx_path = input("Enter path to EOH Index.docx file: ")
    
    # Import the data
    importer = EohIndexImporter()
    result = importer.import_from_docx(docx_path)
    
    if result:
        print("Import successful!")
    else:
        print("Import failed. Please check the file path and try again.")

if __name__ == "__main__":
    main()