from docx import Document

def inspect_document(docx_path):
    print(f"Examining document: {docx_path}")
    doc = Document(docx_path)
    
    print(f"Document has {len(doc.paragraphs)} paragraphs and {len(doc.tables)} tables")
    
    # Sample some content
    if len(doc.tables) > 0:
        table = doc.tables[0]
        rows = len(table.rows)
        cols = len(table.rows[0].cells) if rows > 0 else 0
        print(f"First table has {rows} rows and {cols} columns")
        
        # Show first few rows as sample
        print("\nSample rows:")
        for i in range(min(5, rows)):
            row_text = []
            for cell in table.rows[i].cells:
                cell_text = cell.text.strip()
                if len(cell_text) > 30:
                    cell_text = cell_text[:30] + "..."
                row_text.append(cell_text)
            print(f"Row {i+1}: {row_text}")

if __name__ == "__main__":
    inspect_document("EOH Index.docx")