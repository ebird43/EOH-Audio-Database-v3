import os
import docx

def main():
    print("Simple DOCX File Checker")
    print("=======================")
    
    # List all docx files in the current directory
    docx_files = [f for f in os.listdir('.') if f.endswith('.docx')]
    
    if docx_files:
        print("Found the following .docx files in the current directory:")
        for i, file in enumerate(docx_files, 1):
            print(f"{i}. {file}")
        
        choice = input("\nEnter the number of the file to check (or press Enter to skip): ")
        if choice and choice.isdigit() and 1 <= int(choice) <= len(docx_files):
            file_to_check = docx_files[int(choice)-1]
        else:
            file_to_check = input("Enter the full path to the DOCX file: ")
    else:
        print("No .docx files found in the current directory.")
        file_to_check = input("Enter the full path to the DOCX file: ")
    
    # Remove quotes if present
    file_to_check = file_to_check.strip('"\'')
    
    print(f"\nAttempting to open: {file_to_check}")
    
    try:
        # Try to open the document
        doc = docx.Document(file_to_check)
        
        # If successful, print some basic info
        print(f"Success! Document opened.")
        print(f"Document contains {len(doc.paragraphs)} paragraphs.")
        
        # Print first few paragraphs as a sample
        print("\nFirst 5 paragraphs:")
        for i, para in enumerate(doc.paragraphs[:5], 1):
            text = para.text.strip()
            if text:
                print(f"Paragraph {i}: {text[:50]}..." if len(text) > 50 else f"Paragraph {i}: {text}")
    
    except Exception as e:
        print(f"Error opening document: {str(e)}")
    
    input("\nPress Enter to exit...")

if __name__ == "__main__":
    main()