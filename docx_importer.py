import os
import sqlite3
import re
from datetime import datetime

try:
    from docx import Document
    print("Successfully imported python-docx")
except ImportError:
    print("Error importing docx. Please make sure python-docx is installed.")
    print("Run: pip install python-docx")
    exit(1)

class EohIndexImporter:
    def __init__(self, db_path='adidam_recordings.db'):
        self.db_path = db_path
        self.current_book_id = None
        self.current_book_title = None
        
    def import_from_docx(self, docx_path):
        """Import data from the EOH Index Word document"""
        if not os.path.exists(docx_path):
            print(f"Error: File not found: {docx_path}")
            return False
            
        try:
            # Connect to database
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            # Open the Word document
            print(f"Opening document: {docx_path}")
            doc = Document(docx_path)
            
            print(f"Document opened successfully")
            print(f"Document contains {len(doc.paragraphs)} paragraphs and {len(doc.tables)} tables")
            
            # Extract all text from the document
            full_text = ""
            for para in doc.paragraphs:
                full_text += para.text + "\n"
                
            # For tables - these might contain our data
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text += cell.text + "\n"
            
            print(f"Extracted {len(full_text)} characters of text")
            
            # Begin transaction
            conn.execute('BEGIN TRANSACTION')
            
            # Process the text line by line
            lines = full_text.split('\n')
            line_count = 0
            book_count = 0
            essay_count = 0
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                    
                line_count += 1
                
                # Look for book titles - "The Aletheon" or similar
                book_match = re.search(r'\*\[([^]]+)\]\*', line)
                if book_match:
                    book_title = book_match.group(1).strip()
                    self.process_book_title(book_title, cursor)
                    book_count += 1
                    continue
                
                # Look for essay entries - "**349** Acausal Adidam"
                essay_match = re.search(r'\*\*([^*]+)\*\*\s+(.*)', line)
                if essay_match and self.current_book_id:
                    essay_number = essay_match.group(1).strip()
                    essay_title = essay_match.group(2).strip()
                    self.process_essay_entry(essay_number, essay_title, cursor)
                    essay_count += 1
            
            # Commit the transaction
            conn.commit()
            
            print(f"Processed {line_count} lines")
            print(f"Imported {book_count} books and {essay_count} essays")
            return True
            
        except Exception as e:
            if 'conn' in locals():
                conn.rollback()
            print(f"Error during import: {str(e)}")
            return False
        finally:
            if 'conn' in locals():
                conn.close()
    
    def process_book_title(self, book_title, cursor):
        """Process a book title"""
        # Clean up any remaining formatting
        book_title = book_title.replace("{.smallcaps}", "").strip()
        
        # Check if this book already exists
        cursor.execute("SELECT id FROM books WHERE title = ?", (book_title,))
        existing_book = cursor.fetchone()
        
        if existing_book:
            self.current_book_id = existing_book[0]
            self.current_book_title = book_title
            print(f"Found existing book: {book_title} (ID: {self.current_book_id})")
        else:
            # Insert new book
            cursor.execute(
                "INSERT INTO books (title, display_order) VALUES (?, ?)",
                (book_title, 0)  # We'll update display order later
            )
            self.current_book_id = cursor.lastrowid
            self.current_book_title = book_title
            print(f"Added new book: {book_title} (ID: {self.current_book_id})")
    
    def process_essay_entry(self, essay_number, essay_title, cursor):
        """Process an essay entry"""
        # Clean up title (remove formatting)
        essay_title = re.sub(r'\[([^\]]+)\]\.[\w]+', r'\1', essay_title)
        essay_title = re.sub(r'\{\.[\w]+\}', '', essay_title)
        essay_title = essay_title.replace("[", "").replace("]", "")
        essay_title = essay_title.replace("{.underline}", "").strip()
        
        # Handle multiple recording numbers for the same essay
        essay_numbers = [num.strip() for num in essay_number.split(',')]
        
        for num in essay_numbers:
            # Check if this essay already exists
            cursor.execute(
                "SELECT id FROM essays WHERE title = ? AND book_id = ? AND essay_number = ?", 
                (essay_title, self.current_book_id, num)
            )
            existing_essay = cursor.fetchone()
            
            if existing_essay:
                essay_id = existing_essay[0]
                print(f"  Found existing essay: {essay_title} (ID: {essay_id}, Number: {num})")
            else:
                # Insert new essay
                cursor.execute(
                    """
                    INSERT INTO essays 
                    (title, book_id, essay_number, display_order) 
                    VALUES (?, ?, ?, ?)
                    """,
                    (essay_title, self.current_book_id, num, 0)
                )
                essay_id = cursor.lastrowid
                print(f"  Added new essay: {essay_title} (ID: {essay_id}, Number: {num})")
                
                # For each essay, create a placeholder recording
                cursor.execute(
                    """
                    INSERT INTO recordings 
                    (essay_id, title, date_added) 
                    VALUES (?, ?, ?)
                    """,
                    (essay_id, f"Recording of {essay_title}", datetime.now().isoformat())
                )

def main():
    print("Adidam EOH Index Importer")
    print("=========================")
    
    # Check if the file exists in current directory
    if os.path.exists("EOH Index.docx"):
        docx_path = "EOH Index.docx"
        print(f"Found index file in current directory")
    else:
        docx_path = input("Enter path to EOH Index.docx file: ").strip('"\'')
    
    # Import the data
    importer = EohIndexImporter()
    result = importer.import_from_docx(docx_path)
    
    if result:
        print("Import successful!")
    else:
        print("Import failed. Please check the error messages above.")
    
    input("Press Enter to exit...")

if __name__ == "__main__":
    main()
