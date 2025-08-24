import sqlite3
import os
from datetime import datetime
from docx import Document

class EohTableImporter:
    def __init__(self, db_path='adidam_recordings.db'):
        self.db_path = db_path
        self.current_book_id = None
        self.current_book_title = None
        
    def import_from_docx(self, docx_path):
        """Import data from a Word document with table format"""
        if not os.path.exists(docx_path):
            print(f"Error: File not found: {docx_path}")
            return False
            
        try:
            # Connect to database
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            # Open the Word document
            doc = Document(docx_path)
            print(f"Document opened: {docx_path}")
            
            if len(doc.tables) == 0:
                print("No tables found in document")
                return False
                
            # Process the first table
            table = doc.tables[0]
            rows = len(table.rows)
            print(f"Processing table with {rows} rows")
            
            # Begin transaction
            conn.execute('BEGIN TRANSACTION')
            
            book_count = 0
            essay_count = 0
            
            # Process each row
            for i, row in enumerate(table.rows):
                # Get cell values
                if len(row.cells) < 2:
                    continue
                    
                cell1 = row.cells[0].text.strip()
                cell2 = row.cells[1].text.strip()
                
                if not cell1 or not cell2:
                    continue
                
                # First row or rows with non-numeric first cell might be book titles
                if i == 0 or not cell1[0].isdigit():
                    # This is likely a book title
                    self.process_book_title(cell2, cursor)
                    book_count += 1
                else:
                    # This is likely an essay entry
                    essay_number = cell1
                    essay_title = cell2
                    self.process_essay_entry(essay_number, essay_title, cursor)
                    essay_count += 1
            
            # Commit the transaction
            conn.commit()
            
            print(f"Imported {book_count} books and {essay_count} essays")
            return True
            
        except Exception as e:
            if 'conn' in locals():
                conn.rollback()
            print(f"Error during import: {str(e)}")
            return False
        finally:
            if 'conn' in locals():
                conn.close()
    
    def process_book_title(self, book_title, cursor):
        """Process a book title"""
        # Check if this book already exists
        cursor.execute("SELECT id FROM books WHERE title = ?", (book_title,))
        existing_book = cursor.fetchone()
        
        if existing_book:
            self.current_book_id = existing_book[0]
            self.current_book_title = book_title
            print(f"Found existing book: {book_title} (ID: {self.current_book_id})")
        else:
            # Insert new book
            cursor.execute(
                "INSERT INTO books (title, display_order) VALUES (?, ?)",
                (book_title, 0)  # We'll update display order later
            )
            self.current_book_id = cursor.lastrowid
            self.current_book_title = book_title
            print(f"Added new book: {book_title} (ID: {self.current_book_id})")
    
    def process_essay_entry(self, essay_number, essay_title, cursor):
        """Process an essay entry"""
        if not self.current_book_id:
            print(f"Warning: Essay {essay_number} found without a book. Skipping.")
            return
            
        # Clean special characters
        essay_title = essay_title.replace('\xa0', ' ')
        
        # Handle multiple recording numbers for the same essay
        essay_numbers = [num.strip() for num in essay_number.split(',')]
        
        for num in essay_numbers:
            # Check if this essay already exists
            cursor.execute(
                "SELECT id FROM essays WHERE title = ? AND book_id = ? AND essay_number = ?", 
                (essay_title, self.current_book_id, num)
            )
            existing_essay = cursor.fetchone()
            
            if existing_essay:
                essay_id = existing_essay[0]
                print(f"  Found existing essay: {essay_title} (ID: {essay_id}, Number: {num})")
            else:
                # Insert new essay
                cursor.execute(
                    """
                    INSERT INTO essays 
                    (title, book_id, essay_number, display_order) 
                    VALUES (?, ?, ?, ?)
                    """,
                    (essay_title, self.current_book_id, num, 0)
                )
                essay_id = cursor.lastrowid
                print(f"  Added new essay: {essay_title} (ID: {essay_id}, Number: {num})")
                
                # For each essay, create a placeholder recording
                cursor.execute(
                    """
                    INSERT INTO recordings 
                    (essay_id, title, date_added) 
                    VALUES (?, ?, ?)
                    """,
                    (essay_id, f"Recording of {essay_title}", datetime.now().isoformat())
                )

def main():
    print("Adidam EOH Table Importer")
    print("=========================")
    
    # Check if the file exists in current directory
    if os.path.exists("EOH Index.docx"):
        docx_path = "EOH Index.docx"
        print(f"Found index file in current directory")
    else:
        docx_path = input("Enter path to EOH Index.docx file: ").strip('"\'')
    
    # Import the data
    importer = EohTableImporter()
    result = importer.import_from_docx(docx_path)
    
    if result:
        print("Import successful!")
    else:
        print("Import failed. Please check the error messages above.")
    
    input("Press Enter to exit...")

if __name__ == "__main__":
    main()